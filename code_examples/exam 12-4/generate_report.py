import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = docx.Document()

# Title
title = doc.add_heading('嵌入式系统期末大报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('STM32 独立看门狗(IWDG)与Flash读写操作实验')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)

doc.add_paragraph('\n')

# 1. 实验目的
doc.add_heading('一、 实验目的', level=1)
doc.add_paragraph('1. 深入理解STM32微控制器的独立看门狗(IWDG)的工作原理及配置方法，掌握其在系统异常恢复中的应用。')
doc.add_paragraph('2. 掌握STM32内部Flash的擦除、写入与读取操作，理解Flash存储器的特性及操作注意事项。')
doc.add_paragraph('3. 熟练使用STM32 HAL库进行底层硬件驱动开发。')
doc.add_paragraph('4. 掌握使用PICSimLab等仿真工具进行嵌入式系统软硬件协同仿真的流程。')

# 2. 实验环境
doc.add_heading('二、 实验环境', level=1)
doc.add_paragraph('1. 硬件平台：STM32F103C8T6 (Blue Pill 开发板)')
doc.add_paragraph('2. 软件平台：STM32CubeMX, GCC ARM Embedded Toolchain, Make')
doc.add_paragraph('3. 仿真工具：PICSimLab')

# 3. 实验原理
doc.add_heading('三、 实验原理', level=1)
doc.add_heading('3.1 独立看门狗 (IWDG)', level=2)
doc.add_paragraph('独立看门狗(IWDG)由专用的低速时钟(LSI)驱动，即使主时钟发生故障它也仍然有效。IWDG最适合应用于那些需要看门狗作为一个在主程序之外，能够完全独立工作，并且对时间精度要求较低的场合。当内部向下计数器递减到0时，IWDG会产生一个系统复位信号。在正常运行期间，程序必须定期重载计数器的值（喂狗），以防止系统复位。')

doc.add_heading('3.2 Flash 存储器操作', level=2)
doc.add_paragraph('STM32的内部Flash不仅用于存储程序代码，还可以用于存储掉电不丢失的用户数据。对Flash的写入操作必须在擦除之后进行，且擦除操作以页(Page)为单位。在进行Flash操作时，需要先解锁Flash控制寄存器，操作完成后重新锁定以防止误操作。')

# 4. 关键代码分析
doc.add_heading('四、 关键代码分析', level=1)

doc.add_heading('4.1 IWDG 初始化与喂狗', level=2)
doc.add_paragraph('IWDG的初始化配置了预分频器和重装载值。代码如下：')
code1 = doc.add_paragraph()
code1.add_run('''static void IWDG_Init(void)
{
  IWDG->KR = 0x5555U; // 解锁寄存器访问
  IWDG->PR = 0x04U;   // 预分频器配置
  IWDG->RLR = IWDG_RELOAD_VALUE; // 设置重装载值
  while ((IWDG->SR & (IWDG_SR_PVU | IWDG_SR_RVU)) != 0U) {} // 等待状态更新
  IWDG->KR = 0xAAAAU; // 重载计数器
  IWDG->KR = 0xCCCCU; // 启动IWDG
}

static void IWDG_Feed(void)
{
  IWDG->KR = 0xAAAAU; // 喂狗操作
}''').font.name = 'Courier New'

doc.add_heading('4.2 Flash 擦除与写入', level=2)
doc.add_paragraph('Flash写入前需解锁并擦除目标页，写入过程中需持续喂狗以防超时复位。代码如下：')
code2 = doc.add_paragraph()
code2.add_run('''static AppStatus_t Flash_WriteWords(uint32_t address, const uint32_t *data, uint32_t word_count)
{
  FLASH_EraseInitTypeDef erase = {0}; 
  uint32_t page_error = 0U, i;
  
  if (HAL_FLASH_Unlock() != HAL_OK) return APP_FLASH_WRITE_ERROR;
  
  erase.TypeErase = FLASH_TYPEERASE_PAGES; 
  erase.PageAddress = address; 
  erase.NbPages = 1;
  
  if (HAL_FLASHEx_Erase(&erase, &page_error) != HAL_OK) { 
      HAL_FLASH_Lock(); 
      return APP_FLASH_ERASE_ERROR; 
  }
  
  for (i = 0; i < word_count; i++)
  {
    IWDG_Feed(); // 写入耗时较长，需喂狗
    if (HAL_FLASH_Program(FLASH_TYPEPROGRAM_WORD, address + (i * 4U), data[i]) != HAL_OK)
    { 
        HAL_FLASH_Lock(); 
        return APP_FLASH_WRITE_ERROR; 
    }
  }
  HAL_FLASH_Lock();
  return APP_OK;
}''').font.name = 'Courier New'

doc.add_heading('4.3 错误处理框架', level=2)
doc.add_paragraph('当系统发生严重错误（如Flash读写校验失败）时，进入错误处理函数，停止喂狗并闪烁LED指示错误，最终触发看门狗复位。')
code3 = doc.add_paragraph()
code3.add_run('''void Error_Handler(void)
{
  __disable_irq();
  while (1)
  {
    HAL_GPIO_TogglePin(APP_LED_GPIO_Port, APP_LED_Pin);
    for (volatile uint32_t i = 0; i < 300000U; i++) { }
    // 此处不喂狗，等待IWDG复位系统
  }
}''').font.name = 'Courier New'

# 5. 实验步骤与仿真结果
doc.add_heading('五、 实验步骤与仿真结果', level=1)
doc.add_paragraph('1. 编写并完善 main.c 中的 IWDG 和 Flash 操作代码。')
doc.add_paragraph('2. 使用 Make 工具链编译工程，生成 WatchDoorDog.bin 和 WatchDoorDog.elf 固件文件。')
doc.add_paragraph('3. 配置 PICSimLab 仿真器，选择 Blue_Pill (STM32F103C8T6) 开发板。')
doc.add_paragraph('4. 加载编译生成的固件并运行仿真。')

doc.add_paragraph('仿真运行截图如下：')
if os.path.exists('screenshots/picsimlab_running.png'):
    doc.add_picture('screenshots/picsimlab_running.png', width=Inches(6.0))
elif os.path.exists('screenshots/02_picsimlab_running.png'):
    doc.add_picture('screenshots/02_picsimlab_running.png', width=Inches(6.0))

# 6. 实验总结
doc.add_heading('六、 实验总结', level=1)
doc.add_paragraph('本次实验成功实现了STM32的独立看门狗(IWDG)配置与Flash存储器的读写操作。通过代码编写与PICSimLab仿真，验证了IWDG在系统异常时的复位保护功能，以及Flash在运行时的擦写可靠性。实验过程中，特别注意了在耗时较长的Flash擦写操作中加入喂狗指令，避免了系统的误复位。整体系统运行稳定，达到了预期的设计要求，加深了对嵌入式系统底层硬件操作及容错机制的理解。')

doc.save('report.docx')
print("Report generated successfully.")
